from types import SimpleNamespace

import pytest
from docx import Document

from juridico.revisor_agente import RevisionAgenteError, revisar_documento


class ClienteRevision:
    def __init__(self, resultado):
        self.resultado = resultado
        self.messages = self

    def create(self, **kwargs):
        assert "revision-tutela" in kwargs["system"]
        assert kwargs["tool_choice"]["name"] == "reportar_revision_final"
        bloque = SimpleNamespace(type="tool_use", input=self.resultado)
        return SimpleNamespace(content=[bloque])


def crear_docx(ruta, texto):
    doc = Document()
    doc.add_paragraph(texto)
    doc.save(ruta)


def test_agente_corrige_borrador_sin_cambiar_datos(tmp_path):
    ruta = tmp_path / "tutela.docx"
    crear_docx(ruta, "Ana Pérez solicita que se proceda a la entrega del medicamento.")
    cliente = ClienteRevision({
        "decision": "corregir",
        "resumen_judicial": "Se simplificó una frase.",
        "bloqueos": [],
        "reemplazos": [{
            "anterior": "se proceda a la entrega",
            "nuevo": "entregue",
            "motivo": "Usar un verbo directo.",
        }],
    })

    resultado = revisar_documento(
        "tutela_propia", ruta, {"nombre_agente": "Ana Pérez"}, client=cliente)

    assert resultado["decision"] == "corregir"
    assert "Ana Pérez solicita que entregue el medicamento" in Document(ruta).paragraphs[0].text


def test_agente_bloquea_documento_incompleto(tmp_path):
    ruta = tmp_path / "tutela.docx"
    crear_docx(ruta, "Borrador incompleto")
    cliente = ClienteRevision({
        "decision": "bloquear",
        "resumen_judicial": "Falta el juramento.",
        "bloqueos": [{
            "regla": "32",
            "hallazgo": "Falta el juramento.",
            "dato_faltante": "tutela previa",
            "pregunta_usuario": "¿Ya presentó otra tutela por estos mismos hechos y derechos?",
        }],
        "reemplazos": [],
    })

    with pytest.raises(RevisionAgenteError, match="otra tutela"):
        revisar_documento("tutela_propia", ruta, {}, client=cliente)


def test_agente_no_puede_borrar_dato_confirmado(tmp_path):
    ruta = tmp_path / "tutela.docx"
    crear_docx(ruta, "La accionante es Ana Pérez.")
    cliente = ClienteRevision({
        "decision": "corregir",
        "resumen_judicial": "Cambio inválido.",
        "bloqueos": [],
        "reemplazos": [{
            "anterior": "Ana Pérez",
            "nuevo": "otra persona",
            "motivo": "No permitido.",
        }],
    })

    with pytest.raises(RevisionAgenteError, match="datos confirmados"):
        revisar_documento(
            "tutela_propia", ruta, {"nombre_agente": "Ana Pérez"}, client=cliente)
