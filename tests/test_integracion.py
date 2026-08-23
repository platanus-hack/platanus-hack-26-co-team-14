"""
El recorrido entero, sin WhatsApp y sin Claude.

    texto → extracción → triage → preguntas → minuta → DOCX

Lo que se comprueba es que las piezas encajan: que las preguntas terminan,
que el documento sale, y que cuando no puede salir se dice, en vez de
entregar algo a medias.
"""

from __future__ import annotations

import re

import pytest
from docx import Document

from canal import sesiones
from canal.orquestador import procesar_turno
from tests.doble_claude import ClaudeDoble

TOPE_TURNOS = 30


def conversar(telefono: str, doble: ClaudeDoble, primer_texto: str,
              respuestas: dict[str, str] | None = None):
    """Habla hasta que el sistema deja de preguntar. Devuelve las acciones
    del último turno."""
    respuestas = respuestas or {}
    sesiones.borrar(telefono)
    procesar_turno(telefono, "AUTORIZO", client=doble,
                   mensaje_id=f"consentimiento-{telefono}")

    acciones = procesar_turno(telefono, primer_texto, client=doble)

    for _ in range(TOPE_TURNOS):
        if any(a["tipo"] == "documento" for a in acciones):
            return acciones

        # `esperando` es el dato que el sistema acaba de pedir. Si no pidió
        # nada, la conversación se acabó.
        pendiente = sesiones.obtener(telefono).get("esperando")
        if pendiente is None or pendiente not in respuestas:
            return acciones

        acciones = procesar_turno(telefono, respuestas[pendiente], client=doble)

    pytest.fail(f"la conversación no terminó en {TOPE_TURNOS} turnos; "
                f"último slot preguntado: {doble.preguntado[-3:]}")


def textos(acciones):
    return " ".join(a.get("texto", "") for a in acciones)


# ============================================================
# TUTELA DE UN MENOR — el recorrido completo hasta el DOCX
# ============================================================

def test_tutela_de_un_menor_termina_en_documento(tmp_path):
    doble = ClaudeDoble(
        inicial={
            "tutela_previa_cumplida": False,
            "solicitud_previa": "verbal",
            "sujeto_especial": True,
            "paciente": "menor",
            "eps": "sura",
            "servicio_negado": "el medicamento",
            "diagnostico": "leucemia",
            "edad_menor": "8",
        },
        guion={
            "nombre_completo": "Ana Mosquera Palacios",
            "cedula": "26485912",
            "lugar_expedicion": "Quibdó",
            "ciudad_vulneracion": "Quibdó",
            "direccion_notificaciones": "Calle 5 número 3-20, barrio Yesquita",
            "hecho_vulneracion": ("Llevo tres semanas yendo a la farmacia y "
                                  "siempre me dicen que vuelva mañana"),
            "nombre_menor": "Sara Mosquera",
            "registro_civil_menor": "1098234567",
            "fecha_orden": "12 de marzo",
        },
    )

    acciones = conversar(
        "573001112233", doble,
        "mi hija de 8 años tiene leucemia y la eps sura no le ha entregado "
        "el medicamento, fui a la farmacia y me dijeron que volviera",
        respuestas={k: v for k, v in doble.guion.items()},
    )

    tipos = [a["tipo"] for a in acciones]
    assert "documento" in tipos, textos(acciones)
    assert "audio" in tipos, "la respuesta final tiene que ir también en voz"

    documento = next(a for a in acciones if a["tipo"] == "documento")
    ruta = documento["archivo"]

    doc = Document(ruta)
    contenido = "\n".join(p.text for p in doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                contenido += "\n" + celda.text

    # Ni un placeholder vivo.
    assert not re.search(r"\{\{\w+\}\}", contenido)

    # Los datos de la persona llegaron al papel.
    assert "Ana Mosquera Palacios" in contenido
    assert "Sara Mosquera" in contenido
    assert "26485912" in contenido
    assert "ocho" in contenido          # edad en letras, como pide la minuta
    assert "MEDIDA CAUTELAR PROVISIONAL" not in contenido.upper()
    assert "MEDIDA PROVISIONAL" not in contenido.upper()

    # Y el caso se cerró: no guardamos datos de salud de más.
    assert sesiones.obtener("573001112233")["datos"]["cedula"] is None


# ============================================================
# LO QUE NO PUEDE SALIR, NO SALE
# ============================================================

def test_tutela_propia_termina_en_documento():
    """La minuta propia deriva de la base adulta y no afirma agencia oficiosa."""
    doble = ClaudeDoble(
        inicial={
            "tutela_previa_cumplida": False,
            "solicitud_previa": "verbal",
            "riesgo_vital": True,
            "paciente": "yo",
            "eps": "nueva eps",
            "servicio_negado": "la insulina",
        },
        guion={
            "nombre_completo": "Pedro Nel Rúa",
            "cedula": "71234567",
            "lugar_expedicion": "Medellín",
            "ciudad_vulneracion": "Bogotá",
            "direccion_notificaciones": "Carrera 45 número 12-30",
            "hecho_vulneracion": "me dicen que no hay y que vuelva",
            "diagnostico": "diabetes",
            "fecha_orden": "12 de marzo",
        },
    )

    acciones = conversar(
        "573004445566", doble,
        "no me han dado la insulina y cada día estoy peor, la nueva eps me "
        "dice que vuelva",
        respuestas=dict(doble.guion),
    )

    assert any(a["tipo"] == "documento" for a in acciones), textos(acciones)
    ruta = next(a["archivo"] for a in acciones if a["tipo"] == "documento")
    doc = Document(ruta)
    contenido = "\n".join(p.text for p in doc.paragraphs).lower()
    assert "pedro nel rúa" in contenido
    assert "me dicen que no hay y que vuelva" in contenido
    assert "agente oficioso" not in contenido
    assert "mi madre" not in contenido
    assert "únicamente mientras se decide esta tutela" in contenido
    assert "para evitar que la demora agrave" in contenido
    dicho = textos(acciones).lower()
    assert "@cendoj.ramajudicial.gov.co" in dicho
    assert "el juzgado se encargará de notificarla" in dicho
    audio = next(a["texto"] for a in acciones if a["tipo"] == "audio")
    assert "@" not in audio
    assert "http" not in audio.lower()
    assert "escrito en el chat" in audio.lower()


def test_sin_pedir_nada_va_a_peticion():
    doble = ClaudeDoble(inicial={
        "tutela_previa_cumplida": False,
        "solicitud_previa": "ninguna",
        "eps": "sura",
        "servicio_negado": "una cita con el especialista",
    })

    telefono = "573007778899"
    sesiones.borrar(telefono)
    procesar_turno(telefono, "AUTORIZO", client=doble)
    acciones = procesar_turno(
        telefono,
        "apenas me la formularon y no sé ni a dónde tengo que ir",
        client=doble)
    sesiones.borrar("573007778899")

    dicho = textos(acciones).lower()
    assert "derecho de petición" in dicho
    # Sura está en el catálogo: se da el canal verificado, no uno inventado.
    assert "http" in dicho
    audio = next(a["texto"] for a in acciones if a["tipo"] == "audio")
    assert "http" not in audio.lower()
    assert "escrito en el chat" in audio.lower()


def test_fallo_incumplido_va_a_desacato():
    doble = ClaudeDoble(inicial={"tutela_previa_cumplida": True})
    sesiones.borrar("573001010101")
    procesar_turno("573001010101", "AUTORIZO", client=doble)
    acciones = procesar_turno(
        "573001010101",
        "ya puse una tutela, el juez me dio la razón y la eps no ha cumplido",
        client=doble)
    sesiones.borrar("573001010101")

    assert "desacato" in textos(acciones).lower()


# ============================================================
# EL CANAL
# ============================================================

def test_no_procesa_historia_antes_del_consentimiento():
    telefono = "573002020201"
    doble = ClaudeDoble(inicial={"eps": "sura"})
    sesiones.borrar(telefono)

    acciones = procesar_turno(
        telefono, "Tengo diabetes y la EPS no me entrega la insulina", client=doble)
    caso = sesiones.obtener(telefono)

    assert "autorizo" in textos(acciones).lower()
    assert caso["mensajes"] == []
    assert caso["datos"]["eps"] is None
    assert doble.preguntado == []
    sesiones.borrar(telefono)


def test_consentimiento_explicito_queda_registrado():
    telefono = "573002020200"
    sesiones.borrar(telefono)

    acciones = procesar_turno(
        telefono, "AUTORIZO", mensaje_id="wamid.CONSENTIMIENTO")
    consentimiento = sesiones.obtener(telefono)["consentimiento"]

    assert consentimiento["otorgado"] is True
    assert consentimiento["mensaje_id"] == "wamid.CONSENTIMIENTO"
    assert "muchas gracias" in textos(acciones).lower()
    sesiones.borrar(telefono)

def test_cada_respuesta_va_escrita_y_hablada():
    doble = ClaudeDoble(inicial={})
    acciones = procesar_turno("573002020202", "buenos días", client=doble)
    sesiones.borrar("573002020202")

    tipos = {a["tipo"] for a in acciones}
    assert "audio" in tipos
    assert tipos & {"texto", "botones"}


def test_saludo_solo_invita_a_contar_la_historia_sin_activar_triage():
    telefono = "573002020203"
    doble = ClaudeDoble(inicial={})
    sesiones.borrar(telefono)

    aviso = procesar_turno(telefono, "Hola", client=doble)
    assert "autorizo" in textos(aviso).lower()
    acciones = procesar_turno(telefono, "AUTORIZO", client=doble)
    dicho = textos(acciones).lower()

    assert "cuénteme con sus palabras" in dicho
    assert "tutela" not in dicho
    assert sesiones.obtener(telefono)["esperando"] is None
    assert doble.preguntado == []
    sesiones.borrar(telefono)


def test_triage_empieza_despues_de_recibir_la_historia():
    telefono = "573002020204"
    doble = ClaudeDoble(inicial={})
    sesiones.borrar(telefono)

    procesar_turno(telefono, "AUTORIZO", client=doble)
    acciones = procesar_turno(
        telefono,
        "La EPS no me entrega el medicamento que me ordenó el médico",
        client=doble,
    )

    dicho = textos(acciones).lower()
    assert "antes de seguir" in dicho
    assert "presentado una tutela" in dicho
    assert "cuénteme con sus palabras" not in dicho
    sesiones.borrar(telefono)


def test_reiniciar_borra_el_caso():
    doble = ClaudeDoble(inicial={"eps": "sura"})
    sesiones.borrar("573003030303")
    procesar_turno("573003030303", "AUTORIZO", client=doble)
    procesar_turno("573003030303", "la eps sura no me atiende", client=doble)

    acciones = procesar_turno("573003030303", "empezar de nuevo", client=doble)
    assert sesiones.obtener("573003030303")["datos"]["eps"] is None
    assert "cero" in textos(acciones).lower()
    sesiones.borrar("573003030303")


def test_audio_ininteligible_no_rompe_nada():
    doble = ClaudeDoble(inicial={})
    sesiones.borrar("573005050505")
    procesar_turno("573005050505", "AUTORIZO", client=doble)
    acciones = procesar_turno("573005050505", "", client=doble)
    sesiones.borrar("573005050505")

    assert "repetir" in textos(acciones).lower()


def test_segundo_intento_fallido_pide_respuesta_escrita():
    telefono = "573006060606"
    doble = ClaudeDoble(inicial={})
    sesiones.borrar(telefono)

    procesar_turno(telefono, "AUTORIZO", client=doble)
    procesar_turno(
        telefono,
        "La EPS no me está entregando el medicamento que necesito",
        client=doble,
    )
    primera_repeticion = procesar_turno(telefono, "no se entendió", client=doble)
    segunda_repeticion = procesar_turno(telefono, "tampoco", client=doble)

    assert "intentemos una vez más" in textos(primera_repeticion).lower()
    assert "escríbamela aquí por whatsapp" in textos(segunda_repeticion).lower()
    sesiones.borrar(telefono)


def test_responde_pregunta_lateral_sin_perder_el_dato_pendiente():
    telefono = "573006060607"
    doble = ClaudeDoble(inicial={})
    sesiones.borrar(telefono)

    procesar_turno(telefono, "AUTORIZO", client=doble)
    procesar_turno(
        telefono,
        "La EPS no me está entregando el medicamento que necesito",
        client=doble,
    )
    pendiente_antes = sesiones.obtener(telefono)["esperando"]
    acciones = procesar_turno(
        telefono,
        "¿Hay algún correo electrónico al que pueda enviar esta información?",
        client=doble,
    )

    dicho = textos(acciones).lower()
    assert "cuando terminemos" in dicho
    assert "correo electrónico" in dicho
    assert sesiones.obtener(telefono)["esperando"] == pendiente_antes
    sesiones.borrar(telefono)
