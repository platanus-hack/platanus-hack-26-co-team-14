import re
import sys
from pathlib import Path

from docx import Document

from juridico.mappings.tutela_menor import (
    MAPEO_TUTELA_MENOR,
    REEMPLAZOS_TEXTO_TUTELA_MENOR,
)

from juridico.mappings.tutela_agente import (
    MAPEO_TUTELA_AGENTE,
    REEMPLAZOS_TEXTO_TUTELA_AGENTE,
)


PATRON_X = re.compile(
    r"[Xx]{2,}|(?<=\()[Xx](?=\))"
)


PLANTILLAS = {
    "tutela_menor": {
        "campos":
            MAPEO_TUTELA_MENOR,

        "reglas":
            REEMPLAZOS_TEXTO_TUTELA_MENOR,
    },

    "tutela_agente": {
        "campos":
            MAPEO_TUTELA_AGENTE,

        "reglas":
            REEMPLAZOS_TEXTO_TUTELA_AGENTE,
    },
}


# ============================================================
# RECORRIDO
# ============================================================

def iterar_parrafos(doc):
    for parrafo in doc.paragraphs:
        yield parrafo

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    yield parrafo


# ============================================================
# X -> {{campo}}
# ============================================================

def reemplazar_x_en_parrafo(
    parrafo,
    campos,
    indice,
):
    """
    Cada X consume exactamente una posición del mapping.

    campo == None:
        la X desaparece.

    campo == "eps":
        la X pasa a {{eps}}.
    """

    for run in parrafo.runs:

        texto = run.text

        if not texto:
            continue

        matches = list(
            PATRON_X.finditer(texto)
        )

        if not matches:
            continue

        partes = []
        cursor = 0

        for match in matches:

            if indice >= len(campos):
                raise RuntimeError(
                    "La plantilla contiene más X "
                    "que el mapping. "
                    f"Falló en X #{indice + 1}."
                )

            campo = campos[indice]

            partes.append(
                texto[
                    cursor:
                    match.start()
                ]
            )

            if campo is not None:
                partes.append(
                    "{{"
                    + campo
                    + "}}"
                )

            cursor = match.end()

            indice += 1

        partes.append(
            texto[cursor:]
        )

        run.text = "".join(
            partes
        )

    return indice


# ============================================================
# REGIONES NO-X
# ============================================================

def reemplazar_textos_confirmados(
    parrafo,
    reglas,
    encontrados,
):
    for run in parrafo.runs:

        if not run.text:
            continue

        texto = run.text

        for regla in reglas:

            buscar = regla["buscar"]

            if buscar not in texto:
                continue

            texto = texto.replace(
                buscar,
                regla["reemplazar"],
            )

            encontrados.add(
                regla["nombre"]
            )

        run.text = texto


# ============================================================
# SEMÁNTICOS
# ============================================================

def placeholders_semanticos(doc):
    patron = re.compile(
        r"\{\{([a-zA-Z0-9_]+)\}\}"
    )

    encontrados = []

    for parrafo in iterar_parrafos(doc):

        for match in patron.finditer(
            parrafo.text
        ):

            encontrados.append(
                match.group(1)
            )

    return encontrados


# ============================================================
# PREPARACIÓN
# ============================================================

def preparar_plantilla(
    tipo,
    origen,
    destino,
):

    if tipo not in PLANTILLAS:
        raise ValueError(
            f"Tipo desconocido: {tipo}. "
            f"Disponibles: "
            f"{', '.join(PLANTILLAS.keys())}"
        )

    config = PLANTILLAS[tipo]

    campos = config["campos"]
    reglas = config["reglas"]

    origen = Path(origen)
    destino = Path(destino)

    if not origen.exists():
        raise FileNotFoundError(
            f"No existe: {origen}"
        )

    if origen.suffix.lower() != ".docx":
        raise ValueError(
            "La plantilla debe ser .docx"
        )

    doc = Document(origen)

    parrafos = list(
        iterar_parrafos(doc)
    )

    # ========================================================
    # FASE 1
    # placeholders X
    # ========================================================

    indice = 0

    for parrafo in parrafos:
        indice = reemplazar_x_en_parrafo(
            parrafo,
            campos,
            indice,
        )

    if indice != len(campos):
        raise RuntimeError(
            "Mapping desalineado. "
            f"X encontradas={indice}, "
            f"mapping={len(campos)}"
        )

    # ========================================================
    # FASE 2
    # regiones especiales
    # ========================================================

    reglas_encontradas = set()

    for parrafo in parrafos:
        reemplazar_textos_confirmados(
            parrafo,
            reglas,
            reglas_encontradas,
        )

    nombres_reglas = {
        regla["nombre"]
        for regla in reglas
    }

    faltantes = (
        nombres_reglas
        - reglas_encontradas
    )

    if faltantes:
        raise RuntimeError(
            "No se encontraron regiones "
            "especiales esperadas: "
            + ", ".join(
                sorted(faltantes)
            )
        )

    # ========================================================
    # FASE 3
    # No deben quedar XXXXX
    # ========================================================

    restantes = []

    for parrafo in parrafos:
        if PATRON_X.search(
            parrafo.text
        ):
            restantes.append(
                parrafo.text
            )

    if restantes:
        print()
        print(
            "Párrafos con X restantes:"
        )

        for texto in restantes[:10]:
            print(
                f"  - {texto}"
            )

        raise RuntimeError(
            f"Quedaron {len(restantes)} "
            "párrafos con X."
        )

    # ========================================================
    # FASE 4
    # juez_destino obligatorio
    # ========================================================

    semanticos = placeholders_semanticos(
        doc
    )

    if "juez_destino" not in semanticos:
        raise RuntimeError(
            "La plantilla preparada no contiene "
            "{{juez_destino}}."
        )

    # ========================================================
    # FASE 5
    # guardar
    # ========================================================

    destino.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    doc.save(destino)

    print("=" * 80)
    print(
        "PLANTILLA SEMÁNTICA GENERADA"
    )
    print("=" * 80)

    print(
        f"Tipo                  : {tipo}"
    )

    print(
        f"Origen                : {origen}"
    )

    print(
        f"Destino               : {destino}"
    )

    print(
        f"X procesadas          : {indice}"
    )

    print(
        f"Regiones especiales   : "
        f"{len(reglas_encontradas)}"
    )

    print(
        f"Marcadores semánticos : "
        f"{len(semanticos)}"
    )

    print()

    print(
        "Regiones especiales aplicadas:"
    )

    for nombre in sorted(
        reglas_encontradas
    ):
        print(
            f"  ✓ {nombre}"
        )

    print()

    print(
        "Campos semánticos únicos:"
    )

    for campo in sorted(
        set(semanticos)
    ):
        print(
            f"  - {campo}"
        )


# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":

    if len(sys.argv) != 4:

        print(
            "Uso:\n\n"
            "python -m juridico.preparar_plantilla "
            "<tipo> "
            '"origen.docx" '
            '"destino.docx"\n\n'
            "Tipos:\n"
            "  tutela_menor\n"
            "  tutela_agente"
        )

        sys.exit(1)

    try:

        preparar_plantilla(
            sys.argv[1],
            sys.argv[2],
            sys.argv[3],
        )

    except Exception as e:

        print(
            f"ERROR: {e}"
        )

        sys.exit(1)