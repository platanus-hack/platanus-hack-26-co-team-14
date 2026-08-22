"""
Inspector de plantillas DOCX.

Detecta:
1. Placeholders tipo XXXXX / xxxxx / (X)
2. Texto marcado en rojo aunque no tenga X
3. Clasifica preliminarmente el texto rojo:
   - instruccion
   - alternativa
   - revisar

NO modifica el documento.
"""

import re
import sys
from pathlib import Path

from docx import Document


# Detecta:
# XXXXX
# xxxxx
# (X)
PATRON_X = re.compile(
    r"[Xx]{2,}|(?<=\()[Xx](?=\))"
)


def color_es_rojo(run) -> bool:
    """
    Detecta rojo explícito en el run.

    No usamos esto como única fuente de verdad porque Word
    no siempre expone el color de forma consistente.
    """

    color = run.font.color.rgb

    if color is None:
        return False

    r = color[0]
    g = color[1]
    b = color[2]

    return (
        r >= 160
        and g <= 140
        and b <= 140
        and r > g
        and r > b
    )


def clasificar_texto_rojo(texto: str) -> str:
    """
    Clasificación PRELIMINAR.

    No significa que automáticamente vayamos a reemplazarlo.
    Sirve para encontrar instrucciones escondidas en la plantilla.
    """

    t = texto.lower().strip()

    instrucciones = [
        "escribir",
        "describir",
        "detallar",
        "relacionar",
        "indicar",
        "incluir",
        "diligenciar",
        "colocar",
        "según el caso",
        "dependiendo del caso",
    ]

    if any(x in t for x in instrucciones):
        return "instruccion"

    # Alternativas del tipo:
    # ASIGNAR, O ENTREGAR, O AUTORIZAR, O REALIZAR
    alternativas = [
        "asignar",
        "entregar",
        "autorizar",
        "realizar",
    ]

    encontradas = sum(
        1 for palabra in alternativas
        if palabra in t
    )

    if encontradas >= 2:
        return "alternativa"

    return "revisar"


def texto_rojo_relevante(texto: str) -> bool:
    """
    Evita registrar como candidato cosas como:
        "."
        ","
        "()"

    Debe contener al menos algunas letras.
    """

    letras = re.findall(
        r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]",
        texto
    )

    return len(letras) >= 4


def inspeccionar_parrafo(
    parrafo,
    indice_parrafo,
    origen="documento"
):
    resultados = []

    contexto = parrafo.text

    for indice_run, run in enumerate(parrafo.runs):

        texto = run.text

        if not texto:
            continue

        rojo = color_es_rojo(run)

        # -------------------------------------------------
        # 1. PLACEHOLDERS XXXXX / (X)
        # -------------------------------------------------

        matches = list(
            PATRON_X.finditer(texto)
        )

        for numero_match, match in enumerate(matches):

            resultados.append({
                "tipo": "placeholder_x",

                "id": (
                    f"{origen}:"
                    f"p{indice_parrafo}:"
                    f"r{indice_run}:"
                    f"x{numero_match}"
                ),

                "origen": origen,
                "parrafo": indice_parrafo,
                "run": indice_run,

                "placeholder": match.group(),
                "longitud": len(match.group()),

                "rojo": rojo,
                "clasificacion": "placeholder",

                "texto_run": texto,
                "contexto": contexto,
            })

        # -------------------------------------------------
        # 2. TEXTO ROJO NO REPRESENTADO POR XXXXX
        # -------------------------------------------------

        if rojo:

            # Eliminamos las X para saber si además
            # existe texto real editable en el mismo run.
            texto_sin_x = PATRON_X.sub(
                "",
                texto
            ).strip()

            if (
                texto_sin_x
                and texto_rojo_relevante(texto_sin_x)
            ):

                resultados.append({
                    "tipo": "texto_rojo",

                    "id": (
                        f"{origen}:"
                        f"p{indice_parrafo}:"
                        f"r{indice_run}:red"
                    ),

                    "origen": origen,
                    "parrafo": indice_parrafo,
                    "run": indice_run,

                    "placeholder": texto_sin_x,
                    "longitud": len(texto_sin_x),

                    "rojo": True,

                    "clasificacion":
                        clasificar_texto_rojo(
                            texto_sin_x
                        ),

                    "texto_run": texto,
                    "contexto": contexto,
                })

    return resultados


def inspeccionar_tabla(
    tabla,
    indice_tabla
):
    resultados = []

    for fila_i, fila in enumerate(tabla.rows):

        for celda_i, celda in enumerate(fila.cells):

            for parrafo_i, parrafo in enumerate(
                celda.paragraphs
            ):

                origen = (
                    f"tabla[{indice_tabla}]"
                    f".fila[{fila_i}]"
                    f".celda[{celda_i}]"
                )

                resultados.extend(
                    inspeccionar_parrafo(
                        parrafo,
                        parrafo_i,
                        origen
                    )
                )

    return resultados


def inspeccionar_docx(ruta: str):

    ruta = Path(ruta)

    if not ruta.exists():
        raise FileNotFoundError(
            f"No existe: {ruta}"
        )

    if ruta.suffix.lower() != ".docx":
        raise ValueError(
            "El inspector solo acepta .docx"
        )

    doc = Document(ruta)

    resultados = []

    # Documento principal
    for i, parrafo in enumerate(doc.paragraphs):

        resultados.extend(
            inspeccionar_parrafo(
                parrafo,
                i,
                "documento"
            )
        )

    # Tablas
    for i, tabla in enumerate(doc.tables):

        resultados.extend(
            inspeccionar_tabla(
                tabla,
                i
            )
        )

    return resultados


def mostrar_resultados(
    ruta,
    resultados
):

    placeholders = [
        r for r in resultados
        if r["tipo"] == "placeholder_x"
    ]

    rojos = [
        r for r in resultados
        if r["tipo"] == "texto_rojo"
    ]

    instrucciones = [
        r for r in rojos
        if r["clasificacion"] == "instruccion"
    ]

    alternativas = [
        r for r in rojos
        if r["clasificacion"] == "alternativa"
    ]

    revisar = [
        r for r in rojos
        if r["clasificacion"] == "revisar"
    ]

    print("=" * 100)
    print(f"PLANTILLA: {ruta}")
    print("=" * 100)

    print(
        f"Placeholders X detectados       : "
        f"{len(placeholders)}"
    )

    print(
        f"Textos rojos adicionales        : "
        f"{len(rojos)}"
    )

    print(
        f"  instrucciones                 : "
        f"{len(instrucciones)}"
    )

    print(
        f"  alternativas                  : "
        f"{len(alternativas)}"
    )

    print(
        f"  requieren revisión            : "
        f"{len(revisar)}"
    )

    print()

    for i, r in enumerate(
        resultados,
        start=1
    ):

        print("-" * 100)

        if r["tipo"] == "placeholder_x":

            print(
                f"[{i:03}] PLACEHOLDER_X"
                f" | {r['id']}"
            )

            print(
                f"      valor       : "
                f"{r['placeholder']}"
            )

            print(
                f"      longitud    : "
                f"{r['longitud']}"
            )

            print(
                f"      rojo        : "
                f"{r['rojo']}"
            )

        else:

            print(
                f"[{i:03}] TEXTO_ROJO"
                f" | {r['clasificacion'].upper()}"
                f" | {r['id']}"
            )

            print(
                f'      editable?   : '
                f'"{r["placeholder"]}"'
            )

        print(
            f"      origen      : "
            f"{r['origen']}"
        )

        print(
            f"      párrafo     : "
            f"{r['parrafo']}"
        )

        print(
            f"      run         : "
            f"{r['run']}"
        )

        print(
            f'      contexto    : '
            f'"{r["contexto"]}"'
        )

    print("-" * 100)


if __name__ == "__main__":

    if len(sys.argv) != 2:

        print(
            "Uso:\n"
            "python juridico/"
            "inspeccionar_plantilla.py "
            '"archivo.docx"'
        )

        sys.exit(1)

    archivo = sys.argv[1]

    try:

        resultados = inspeccionar_docx(
            archivo
        )

        mostrar_resultados(
            archivo,
            resultados
        )

    except Exception as e:

        print(
            f"ERROR: {e}"
        )

        sys.exit(1)