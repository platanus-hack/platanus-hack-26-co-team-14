"""
Render determinístico de documentos jurídicos.

Responsabilidad:
    plantilla preparada + datos confirmados -> DOCX final

Este módulo:
- NO llama LLM.
- NO pregunta datos.
- NO decide la ruta jurídica.
- NO selecciona jurisprudencia.
- NO selecciona juzgados.
- NO inventa información.

Los datos deben llegar ya validados desde el flujo conversacional.
"""

import re
from copy import deepcopy
from pathlib import Path

from docx import Document

from datos.juzgados import contexto_juzgado


# ============================================================
# RUTAS
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

PLANTILLAS_DIR = BASE_DIR / "plantillas"


PLANTILLAS = {
    "tutela_propia": (
        PLANTILLAS_DIR
        / "tutela_propia_preparada.docx"
    ),
    "tutela_menor": (
        PLANTILLAS_DIR
        / "tutela_menor_preparada.docx"
    ),

    "tutela_agente": (
        PLANTILLAS_DIR
        / "tutela_agente_preparada.docx"
    ),
}


# {{campo}}
PATRON_CAMPO = re.compile(
    r"\{\{([a-zA-Z0-9_]+)\}\}"
)


# ============================================================
# UTILIDADES DOCX
# ============================================================

def iterar_parrafos(doc):
    """
    Recorre párrafos del documento principal y tablas.
    """

    for parrafo in doc.paragraphs:
        yield parrafo

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    yield parrafo


def extraer_campos_plantilla(doc):
    """
    Obtiene todos los campos {{...}} existentes
    en una plantilla preparada.
    """

    campos = []

    for parrafo in iterar_parrafos(doc):
        for match in PATRON_CAMPO.finditer(
            parrafo.text
        ):
            campos.append(
                match.group(1)
            )

    return campos


# ============================================================
# VALORES
# ============================================================

def convertir_valor(valor):
    """
    Todo termina convertido a texto.

    None NO se convierte silenciosamente.
    """

    if valor is None:
        return None

    if isinstance(valor, bool):
        return (
            "Sí"
            if valor
            else "No"
        )

    return str(valor).strip()


def datos_limpios(datos: dict):
    """
    Limpia los datos sin inventar valores.
    """

    salida = {}

    for clave, valor in datos.items():
        salida[clave] = convertir_valor(
            valor
        )

    return salida


# ============================================================
# VALIDACIÓN
# ============================================================

def validar_datos(
    doc,
    datos: dict,
):
    """
    Verifica que todos los campos requeridos por la
    plantilla existan y tengan valor.
    """

    campos = set(
        extraer_campos_plantilla(doc)
    )

    faltantes = []

    for campo in sorted(campos):

        if campo not in datos:
            faltantes.append(campo)
            continue

        valor = datos[campo]

        if valor is None:
            faltantes.append(campo)
            continue

        if (
            isinstance(valor, str)
            and not valor.strip()
        ):
            faltantes.append(campo)

    return faltantes


# ============================================================
# REEMPLAZO
# ============================================================

def reemplazar_campos_en_parrafo(
    parrafo,
    datos: dict,
):
    """
    Reemplaza placeholders dentro de los runs.

    Las plantillas preparadas por preparar_plantilla.py
    dejan los {{campos}} dentro de runs individuales,
    por lo que mantenemos el formato original.
    """

    reemplazados = []

    for run in parrafo.runs:

        texto = run.text

        if not texto:
            continue

        matches = list(
            PATRON_CAMPO.finditer(texto)
        )

        if not matches:
            continue

        nuevo = texto

        # set evita reemplazar el mismo campo varias
        # veces dentro del mismo run innecesariamente.
        campos_run = {
            match.group(1)
            for match in matches
        }

        for campo in campos_run:

            if campo not in datos:
                continue

            valor = datos[campo]

            if valor is None:
                continue

            marcador = (
                "{{"
                + campo
                + "}}"
            )

            nuevo = nuevo.replace(
                marcador,
                str(valor),
            )

            reemplazados.append(
                campo
            )

        run.text = nuevo

    return reemplazados


def reemplazar_campos(
    doc,
    datos,
):
    reemplazados = []

    for parrafo in iterar_parrafos(doc):

        reemplazados.extend(
            reemplazar_campos_en_parrafo(
                parrafo,
                datos,
            )
        )

    return reemplazados


# ============================================================
# VALIDACIÓN POST-RENDER
# ============================================================

def placeholders_restantes(doc):
    """
    Ningún {{campo}} debería sobrevivir al render final.
    """

    restantes = []

    for parrafo in iterar_parrafos(doc):

        matches = list(
            PATRON_CAMPO.finditer(
                parrafo.text
            )
        )

        for match in matches:

            restantes.append({
                "campo": match.group(1),
                "contexto": parrafo.text,
            })

    return restantes


# ============================================================
# ENRIQUECIMIENTO TERRITORIAL
# ============================================================

def agregar_juzgado(
    datos: dict,
):
    """
    Si tenemos ciudad_vulneracion, resolvemos el
    contexto territorial.

    datos/juzgados.py es quien toma esa decisión.

    render.py solamente incorpora el resultado.
    """

    salida = deepcopy(datos)

    ciudad = salida.get(
        "ciudad_vulneracion"
    )

    departamento = salida.get(
        "departamento_vulneracion"
    )

    if not ciudad:
        return salida

    juzgado = contexto_juzgado(
        ciudad=ciudad,
        departamento=departamento,
    )

    salida.update(
        juzgado
    )

    return salida


# ============================================================
# RENDER PRINCIPAL
# ============================================================

def renderizar_documento(
    tipo: str,
    datos: dict,
    salida,
    resolver_juzgado=True,
):
    """
    Renderiza un documento jurídico.

    Ejemplo:

        renderizar_documento(
            "tutela_menor",
            datos,
            "salidas/caso_001.docx"
        )
    """

    if tipo not in PLANTILLAS:
        raise ValueError(
            f"Tipo desconocido: {tipo}. "
            f"Disponibles: "
            f"{', '.join(PLANTILLAS.keys())}"
        )

    plantilla = PLANTILLAS[tipo]

    if not plantilla.exists():

        raise FileNotFoundError(
            "No existe la plantilla preparada: "
            f"{plantilla}"
        )

    salida = Path(salida)

    # --------------------------------------------------------
    # 1. Cargar plantilla preparada
    # --------------------------------------------------------

    doc = Document(
        plantilla
    )

    # --------------------------------------------------------
    # 2. Copiar datos
    # --------------------------------------------------------

    contexto = deepcopy(
        datos
    )

    # --------------------------------------------------------
    # 3. Juzgado
    # --------------------------------------------------------

    if resolver_juzgado:

        contexto = agregar_juzgado(
            contexto
        )

    # --------------------------------------------------------
    # 4. Normalizar
    # --------------------------------------------------------

    contexto = datos_limpios(
        contexto
    )

    # --------------------------------------------------------
    # 5. Validar requisitos de plantilla
    # --------------------------------------------------------

    faltantes = validar_datos(
        doc,
        contexto,
    )

    if faltantes:

        raise ValueError(
            "No se puede generar el documento. "
            "Faltan datos requeridos:\n"
            + "\n".join(
                f"  - {campo}"
                for campo in faltantes
            )
        )

    # --------------------------------------------------------
    # 6. Reemplazar
    # --------------------------------------------------------

    reemplazados = reemplazar_campos(
        doc,
        contexto,
    )

    # --------------------------------------------------------
    # 7. Comprobar que no quede {{...}}
    # --------------------------------------------------------

    restantes = placeholders_restantes(
        doc
    )

    if restantes:

        detalles = "\n".join(
            (
                f"  - {r['campo']}: "
                f"{r['contexto']}"
            )
            for r in restantes
        )

        raise RuntimeError(
            "El documento todavía contiene "
            "placeholders sin resolver:\n"
            + detalles
        )

    # --------------------------------------------------------
    # 8. Crear directorio
    # --------------------------------------------------------

    salida.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    # --------------------------------------------------------
    # 9. Guardar
    # --------------------------------------------------------

    doc.save(
        salida
    )

    # --------------------------------------------------------
    # 10. Resultado estructurado
    # --------------------------------------------------------

    return {
        "ok": True,

        "tipo":
            tipo,

        "archivo":
            str(salida),

        "campos_reemplazados":
            sorted(
                set(reemplazados)
            ),

        "cantidad_reemplazos":
            len(reemplazados),

        "juez_destino":
            contexto.get(
                "juez_destino"
            ),

        "email_juzgado":
            contexto.get(
                "email_juzgado"
            ),

        "dane_juzgado":
            contexto.get(
                "dane_juzgado"
            ),

        "requiere_revision_juzgado":
            contexto.get(
                "requiere_revision_juzgado"
            ),
    }
